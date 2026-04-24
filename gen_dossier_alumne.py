from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Marges A4 ──────────────────────────────────────────────────
for section in doc.sections:
    section.page_height = Cm(29.7)
    section.page_width  = Cm(21.0)
    section.top_margin    = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Colors ─────────────────────────────────────────────────────
TARONJA  = RGBColor(0xF5, 0xA6, 0x23)
VERMELL  = RGBColor(0xE6, 0x39, 0x46)
GRIS     = RGBColor(0x44, 0x4B, 0x6A)
GRISC    = RGBColor(0x88, 0x8E, 0xA8)
NEGRE    = RGBColor(0x1A, 0x1A, 0x2E)
VERD     = RGBColor(0x52, 0xB7, 0x88)

def set_color(run, color):
    run.font.color.rgb = color

def add_page_break(doc):
    doc.add_page_break()

def heading(doc, text, level=1, color=NEGRE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16 if level == 1 else 13 if level == 2 else 11)
    run.font.color.rgb = color
    return p

def subheading(doc, text, color=TARONJA):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(7.5)
    run.font.color.rgb = color
    # Espaiat de lletres simulat amb estil
    return p

def body(doc, text, size=10, color=NEGRE, italic=False, bold=False, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.italic = italic
    run.bold = bold
    return p

def question_block(doc, num, total, text, scaffold_rows, bonus=False, obligatori=False, extra_lines=6):
    """Bloc de pregunta amb scaffolding i espai de resposta."""
    # Capçalera
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    if bonus:
        r = p.add_run("⭐  Fast finisher — opcional")
        r.bold = True; r.font.size = Pt(8); r.font.color.rgb = TARONJA
    elif obligatori:
        r1 = p.add_run(f"Pregunta {num} de {total}   ")
        r1.font.size = Pt(8); r1.font.color.rgb = GRISC
        r2 = p.add_run("· OBLIGATÒRIA")
        r2.bold = True; r2.font.size = Pt(8); r2.font.color.rgb = VERMELL
    else:
        r = p.add_run(f"Pregunta {num} de {total}")
        r.font.size = Pt(8); r.font.color.rgb = GRISC

    # Text de la pregunta
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(4)
    run = p2.add_run(text)
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = NEGRE

    # Scaffolding (requadre)
    if scaffold_rows:
        # Títol scaffolding
        ps = doc.add_paragraph()
        ps.paragraph_format.space_after = Pt(1)
        ps.paragraph_format.left_indent = Cm(0.3)
        rs = ps.add_run("ESTRUCTURA DE RESPOSTA")
        rs.bold = True; rs.font.size = Pt(7); rs.font.color.rgb = TARONJA

        # Files scaffolding
        for label in scaffold_rows:
            pr = doc.add_paragraph()
            pr.paragraph_format.space_after = Pt(1)
            pr.paragraph_format.left_indent = Cm(0.3)
            r1 = pr.add_run(f"{label}  ")
            r1.font.size = Pt(9); r1.font.color.rgb = TARONJA; r1.bold = True
            r2 = pr.add_run("_" * 55)
            r2.font.size = Pt(9); r2.font.color.rgb = GRISC

        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # Línies de resposta
    for _ in range(extra_lines):
        pl = doc.add_paragraph()
        pl.paragraph_format.space_after = Pt(1)
        rl = pl.add_run("_" * 90)
        rl.font.size = Pt(9)
        rl.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

def session_header(doc, missio, sessio, titol, color=VERMELL):
    add_page_break(doc)
    # Pastilla missió
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"  {missio}  ")
    r.bold = True; r.font.size = Pt(8); r.font.color.rgb = color
    # Títol sessió
    heading(doc, f"{sessio} — {titol}", level=2, color=color)
    # Línia separadora
    p = doc.add_paragraph("─" * 80)
    p.paragraph_format.space_after = Pt(6)
    p.runs[0].font.size = Pt(7)
    p.runs[0].font.color.rgb = RGBColor(0xDD, 0xDD, 0xDD)


# ════════════════════════════════════════════════════════
#  PORTADA
# ════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(60)
p.paragraph_format.space_after  = Pt(4)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("⚡  CRISI ENERGÈTICA")
r.bold = True; r.font.size = Pt(28); r.font.color.rgb = TARONJA

p2 = doc.add_paragraph("Consell Assessor Jove · Unió Europea")
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(2)
p2.runs[0].font.size = Pt(12); p2.runs[0].font.color.rgb = GRISC

p3 = doc.add_paragraph("Institut de Matadepera · Curs 2025-26")
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_after = Pt(50)
p3.runs[0].font.size = Pt(10); p3.runs[0].font.color.rgb = GRISC

# Línia avís simulació
pa = doc.add_paragraph()
pa.alignment = WD_ALIGN_PARAGRAPH.CENTER
pa.paragraph_format.space_after = Pt(60)
ra = pa.add_run("🎓  Simulació educativa basada en fets reals. Les dades han estat adaptades amb fins pedagògics.")
ra.italic = True; ra.font.size = Pt(9); ra.font.color.rgb = RGBColor(0x88, 0x70, 0x30)

# Dades alumne
for label in ["Nom i cognoms:", "Grup:", "Data d'inici:"]:
    pf = doc.add_paragraph()
    pf.paragraph_format.space_after = Pt(14)
    r1 = pf.add_run(f"{label}   ")
    r1.bold = True; r1.font.size = Pt(11); r1.font.color.rgb = NEGRE
    r2 = pf.add_run("_" * 45)
    r2.font.size = Pt(11); r2.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)


# ════════════════════════════════════════════════════════
#  SESSIÓ 1 — Diagnosi inicial
# ════════════════════════════════════════════════════════
session_header(doc, "MISSIÓ 1", "Sessió 1", "Diagnosi inicial", VERMELL)

body(doc, "Respon individualment. Podràs comparar aquestes respostes amb les de la Sessió 7.", italic=True, color=GRISC)

question_block(doc, 1, 2,
    "Quan sents «crisi energètica», de què creus que parla?",
    ["Crec que tracta de…", "Un exemple que conec:", "Una cosa que em genera dubtes:"],
    extra_lines=5)

question_block(doc, 2, 2,
    "Creus que t'afecta? Com?",
    ["Sí / No, perquè…", "Un exemple concret:", "Algú que crec que s'afecta més:"],
    extra_lines=5)

question_block(doc, None, None,
    "Qui creus que té la culpa? (intuïció, no cal justificar)",
    None, bonus=True, extra_lines=3)


# ════════════════════════════════════════════════════════
#  SESSIÓ 2 — Analitzar fonts
# ════════════════════════════════════════════════════════
session_header(doc, "MISSIÓ 1", "Sessió 2", "A qui creure? Anàlisi de fonts", VERMELL)

body(doc, "Sobre que has rebut: ______   (A · B · C · D — encercla'l)", bold=True)
body(doc, "Respon a partir del document del teu sobre.", italic=True, color=GRISC)

question_block(doc, 1, 2,
    "Qui ha escrit aquest document i per quins motius? Quins interessos pot tenir l'organització o la persona que l'ha redactat?",
    ["Autor/organització:", "Motiu probable:", "Indici que m'ho fa pensar:"],
    extra_lines=5)

question_block(doc, 2, 2,
    "Hi ha alguna dada o argument al document que et sembla sòlid i ben fonamentat? I algun que et genera dubtes o que podria estar esbiaixat? Explica el per quès.",
    ["Sòlid perquè…", "Dubtós perquè…", "Possible biaix detectat:"],
    extra_lines=6)

question_block(doc, None, None,
    "Quina part creus que és fiable i per quins criteris ho decideixes? Pots imaginar un lector que s'ho cregués tot sense qüestionar-ho? Qui seria?",
    None, bonus=True, extra_lines=4)


# ════════════════════════════════════════════════════════
#  SESSIÓ 3 — Poder i decisió
# ════════════════════════════════════════════════════════
session_header(doc, "MISSIÓ 2", "Sessió 3", "Qui decideix? Poder i decisió",
               RGBColor(0xE8, 0x7D, 0x0D))

body(doc, "Respon tenint en compte els cinc actors (govern, empresa energètica, ONG, acadèmia, ciutadania) i les pistes de l'obertura.", italic=True, color=GRISC)

question_block(doc, 1, 3,
    "Qui creus que té més poder REAL en la decisió energètica europea? Per quins motius concrets?",
    ["L'actor amb més poder:", "Una raó concreta:", "Prova o indici:"],
    extra_lines=5)

question_block(doc, 2, 3,
    "Qui creus que HAURIA DE TENIR més poder per decidir sobre l'energia? Per quins motius?",
    ["Hauria de tenir poder:", "Perquè (valor/principi):", "Una objecció possible:"],
    extra_lines=5)

question_block(doc, 3, 3,
    "Com a Consell Assessor, a qui escoltaríeu per elaborar la vostra proposta? Podeu donar veu a uns i no als altres?",
    ["Escoltem:", "No escoltem:", "Criteri per decidir-ho:"],
    extra_lines=5)

question_block(doc, None, None,
    "Hi ha diferència entre les teves dues primeres respostes? Si és que sí, què ens diu això sobre com funciona la democràcia real?",
    None, bonus=True, extra_lines=4)


# ════════════════════════════════════════════════════════
#  SESSIÓ 4 — Les opcions energètiques
# ════════════════════════════════════════════════════════
session_header(doc, "MISSIÓ 2", "Sessió 4", "El mapa del poder — Les opcions energètiques",
               RGBColor(0xE8, 0x7D, 0x0D))

body(doc, "Sobre que has rebut: ______   (A · B · C · D — encercla'l)", bold=True)
body(doc, "Respon a partir de la fitxa de la teva opció.", italic=True, color=GRISC)

question_block(doc, 1, 2,
    "Quins riscos veus en l'opció que has analitzat? Hi ha riscos que no apareixen explícitament a la fitxa però que intuïtes?",
    ["Risc principal (de la fitxa):", "Raó que l'avala:", "Risc que no apareix a la fitxa:"],
    extra_lines=5)

question_block(doc, 2, 2,
    "Si haguessis de descartar una de les tres opcions, quina seria i per quins motius? Pots descartar-ne la teva pròpia si creus que té massa riscos.",
    ["Descartaria:", "Motiu principal:", "Reconec que té el valor de:"],
    extra_lines=5)

question_block(doc, None, None,
    "Hi ha una Opció D que no surt al mapa? Basant-te en les pistes de l'obertura, formula-la i argumenta per quins motius no s'ha considerat fins ara.",
    None, bonus=True, extra_lines=5)


# ════════════════════════════════════════════════════════
#  SESSIÓ 5 — Negociació
# ════════════════════════════════════════════════════════
session_header(doc, "MISSIÓ 3", "Sessió 5", "Negociar o trencar",
               RGBColor(0xD4, 0xAC, 0x0D))

body(doc, "Sobre que has rebut: ______   (R · G · X · E — encercla'l)", bold=True)
body(doc, "Respon a partir del document del teu sobre de negociació.", italic=True, color=GRISC)

question_block(doc, 1, 2,
    "Quins interessos té el teu actor? Què vol REALMENT més enllà de l'energia — poder, reconeixement, seguretat, diners, prestigi?",
    ["Vol explícitament:", "Vol realment (motivació fonda):", "Línia vermella (mai acceptarà):"],
    extra_lines=6)

question_block(doc, 2, 2,
    "Què podria oferir Europa al teu actor que NO sigui energia ni diners directament? Pensa en garanties, reconeixements, acords comercials, estatus internacional...",
    ["Oferta possible 1:", "Oferta possible 2 (no és diners):", "Per quins motius li interessaria:"],
    extra_lines=5)

question_block(doc, None, None,
    "Hi ha alguna cosa que Europa i el teu actor tinguin en comú? Quin podria ser el terreny compartit sobre el qual construir l'acord?",
    None, bonus=True, extra_lines=4)


# ════════════════════════════════════════════════════════
#  SESSIÓ 6 — La proposta final
# ════════════════════════════════════════════════════════
session_header(doc, "MISSIÓ 4", "Sessió 6", "La nostra proposta", VERD)

body(doc, "Reflexió individual. La pregunta 3 és obligatòria per a tothom.", italic=True, color=GRISC)

question_block(doc, 1, 3,
    "Quina ha estat la idea, pregunta o descoberta que més t'ha fet pensar durant tot el projecte? Per quins motius?",
    ["La idea/descoberta és:", "M'ha fet pensar perquè:", "Ha canviat la meva visió en:"],
    extra_lines=5)

question_block(doc, 2, 3,
    "Quina decisió creus que Europa hauria de prendre sobre l'energia? Per quins motius, i quins valors guien la teva resposta?",
    ["La decisió hauria de ser:", "L'argumento amb:", "El valor que guia la resposta:"],
    extra_lines=6)

question_block(doc, 3, 3,
    "Qui podria sortir perjudicat per la decisió que defenses? Treballadors d'indústries fòssils, països productors, comunitats mineres, generacions futures... És acceptable aquest cost? Per quins motius?",
    ["Perjudicats concrets:", "Cost que hauran d'assumir:", "És acceptable perquè / malgrat que:"],
    obligatori=True, extra_lines=7)


# ════════════════════════════════════════════════════════
#  SESSIÓ 7 — Reflexió final
# ════════════════════════════════════════════════════════
session_header(doc, "MISSIÓ 4", "Sessió 7", "L'Assemblea i reflexió final", VERD)

body(doc, "Recupera la fitxa de la Sessió 1 i compara el que pensaves llavors amb el que penses ara.", italic=True, color=GRISC)

question_block(doc, 1, 3,
    "Què pensaves al principi sobre la crisi energètica que ara has canviat o matisat? Explica concretament el canvi.",
    ["Abans pensava que…", "Ara penso que…", "El que em va fer canviar va ser:"],
    extra_lines=6)

question_block(doc, 2, 3,
    "Quina ha estat la pregunta més difícil de respondre durant tot el projecte? Per quins motius et va costar tant?",
    ["La pregunta més difícil:", "Em va costar perquè…", "Ara la respondria diferent?"],
    extra_lines=5)

question_block(doc, 3, 3,
    "Com a ciutadà o ciutadana, QUÈ POTS FER TU amb tot el que has après? Pensa en accions concretes i realistes, no en grans declaracions.",
    None,
    extra_lines=7)


# ════════════════════════════════════════════════════════
#  CONTRACTE DE VALORS
# ════════════════════════════════════════════════════════
add_page_break(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
p.paragraph_format.space_after  = Pt(6)
r = p.add_run("EL MEU CONTRACTE DE VALORS")
r.bold = True; r.font.size = Pt(14); r.font.color.rgb = TARONJA

p2 = doc.add_paragraph("El Consell Assessor Jove ha acabat la seva feina. Ara et toca a tu.")
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(20)
p2.runs[0].italic = True; p2.runs[0].font.size = Pt(10); p2.runs[0].font.color.rgb = GRISC

compromisos = [
    ("① Com a consumidor/a d'energia,", "em comprometo a…"),
    ("② Com a ciutadà/ana que s'informa i detecta el biaix,", "em comprometo a…"),
    ("③ Com a persona amb responsabilitat global,", "em comprometo a…"),
]

for titol, subtitol in compromisos:
    pc = doc.add_paragraph()
    pc.paragraph_format.space_before = Pt(12)
    pc.paragraph_format.space_after  = Pt(3)
    r1 = pc.add_run(titol + " ")
    r1.bold = True; r1.font.size = Pt(10.5); r1.font.color.rgb = TARONJA
    r2 = pc.add_run(subtitol)
    r2.font.size = Pt(10.5); r2.font.color.rgb = NEGRE

    for _ in range(4):
        pl = doc.add_paragraph()
        pl.paragraph_format.space_after = Pt(1)
        pl.runs  # buit per tenir la línia
        rl = pl.add_run("_" * 90)
        rl.font.size = Pt(10); rl.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

# Signatura
doc.add_paragraph()
pf = doc.add_paragraph()
pf.paragraph_format.space_before = Pt(30)
r1 = pf.add_run("Signatura:   ")
r1.bold = True; r1.font.size = Pt(10)
r2 = pf.add_run("_" * 35)
r2.font.size = Pt(10); r2.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
r3 = pf.add_run("      Data:   ")
r3.bold = True; r3.font.size = Pt(10)
r4 = pf.add_run("_" * 20)
r4.font.size = Pt(10); r4.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

pn = doc.add_paragraph()
pn.paragraph_format.space_before = Pt(20)
pn.alignment = WD_ALIGN_PARAGRAPH.CENTER
rn = pn.add_run("No els compartiràs si no vols. Són teus.")
rn.italic = True; rn.font.size = Pt(9); rn.font.color.rgb = GRISC


# ════════════════════════════════════════════════════════
#  DESAR
# ════════════════════════════════════════════════════════
out = "/home/user/crisi-energetica/dossier_alumne.docx"
doc.save(out)
print(f"Desat: {out}")
